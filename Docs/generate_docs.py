import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def format_paragraph(p, space_after=6, space_before=0, line_spacing=1.15):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    format_paragraph(h, space_before=16, space_after=8)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x4D, 0x2B) # Forest Green
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    format_paragraph(h, space_before=12, space_after=6)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50) # Dark Slate
    return h

def add_heading_3(doc, text):
    h = doc.add_paragraph()
    format_paragraph(h, space_before=10, space_after=4)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    return h

def add_body_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    format_paragraph(p, space_after=6)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    format_paragraph(p, space_after=4)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x1E, 0x4D, 0x2B)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def add_code_block(doc, code_text, title=None):
    if title:
        p_t = doc.add_paragraph()
        format_paragraph(p_t, space_before=6, space_after=2)
        r_t = p_t.add_run(title)
        r_t.font.name = 'Calibri'
        r_t.font.size = Pt(10)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4F6F7")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>
            <w:left w:val="single" w:sz="18" w:space="0" w:color="1E4D2B"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    doc.add_paragraph() # Spacer

def style_table_header(row, bg_hex="1E4D2B"):
    for cell in row.cells:
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            format_paragraph(p, space_before=2, space_after=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def style_table_body(table, col_widths=None):
    set_table_borders(table)
    for i, row in enumerate(table.rows[1:]):
        bg = "F9FAFA" if i % 2 == 1 else "FFFFFF"
        for j, cell in enumerate(row.cells):
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if col_widths and j < len(col_widths):
                cell.width = Inches(col_widths[j])
            for p in cell.paragraphs:
                format_paragraph(p, space_before=2, space_after=2)
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def generate_document():
    doc = docx.Document()
    
    # Set page margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- TITLE PAGE / COVER ---
    title_p = doc.add_paragraph()
    format_paragraph(title_p, space_before=40, space_after=12)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run("AgriVision")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(36)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1E, 0x4D, 0x2B)

    sub_p = doc.add_paragraph()
    format_paragraph(sub_p, space_before=0, space_after=24)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Smart Precision Agriculture & Generative AI Platform\nComplete Software Architecture & System Documentation")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(16)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_div = divider.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    r_div.font.color.rgb = RGBColor(0x1E, 0x4D, 0x2B)

    meta_p = doc.add_paragraph()
    format_paragraph(meta_p, space_before=40, space_after=40)
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = meta_p.add_run("System Specification & Architecture Documentation\nVersion 2.0 | July 2026\nPrepared for Project Review & Portfolio")
    r_meta.font.name = 'Calibri'
    r_meta.font.size = Pt(11)
    r_meta.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # --- TABLE OF CONTENTS SUMMARY ---
    add_heading_1(doc, "Document Table of Contents")
    add_bullet_p(doc, "1. Scope of the Project")
    add_bullet_p(doc, "2. Functional Requirements & Non-Functional Requirements")
    add_bullet_p(doc, "3. Use Case Diagram")
    add_bullet_p(doc, "4. Adopted Methodology")
    add_bullet_p(doc, "5. Work Plan (MS Project Schedule & WBS)")
    add_bullet_p(doc, "6. Entity Relationship Diagram (ERD)")
    add_bullet_p(doc, "7. Architecture Design Diagram")
    add_bullet_p(doc, "8. Sequence Diagrams & Usage Scenarios")
    add_bullet_p(doc, "9. Class Diagram")
    add_bullet_p(doc, "10. Database Design & Data Dictionary")

    doc.add_paragraph()

    # ==========================================
    # SECTION 1: SCOPE OF THE PROJECT
    # ==========================================
    add_heading_1(doc, "1. Scope of the Project")
    
    add_heading_2(doc, "1.1 System Background & Industry Context")
    add_body_p(doc, "Modern agriculture faces critical challenges including unpredictable climate shifts, soil degradation, inefficient water management, and pests. Traditional farming relies heavily on manual observation, which often leads to delayed interventions and suboptimal crop yields. Smallholder and commercial farmers in agricultural regions (such as Punjab, Pakistan) require real-time field data, satellite remote sensing, and intelligent localized agronomic recommendations to maximize crop output and ensure resource efficiency.")
    add_body_p(doc, "AgriVision is an end-to-end Smart Precision Agriculture and Generative AI advisory platform designed to bridge this gap. The platform integrates internet-of-things (IoT) field hardware, multi-spectral satellite remote sensing, spatial geospatial database modeling, and Generative Artificial Intelligence (Google Gemini LLM) into a unified native mobile application for farmers and agronomists.")

    add_heading_2(doc, "1.2 Platform Architectural Scope")
    add_body_p(doc, "The scope of the AgriVision project spans four major technical domains:")
    
    add_bullet_p(doc, "Micro-controller node equipped with soil moisture sensors, soil temperature sensors (Dallas 1-Wire DS18B20), air temperature/humidity sensors, and Modbus RS485 soil NPK (Nitrogen, Phosphorus, Potassium) sensors. Features dual-mode operation: WiFi MQTT publishing in production mode and USB Serial JSON payload streaming in development mode, complete with Arduino Over-The-Air (OTA) firmware updates.", "1. Hardware IoT Field Node (ESP32): ")
    
    add_bullet_p(doc, "FastAPI (Python 3.11+) asynchronous microservice engine operating on Uvicorn/Gunicorn. Manages spatial geospatial field boundaries (PostGIS SRID 4326), high-frequency time-series telemetry hypertables (TimescaleDB), background synchronization workers, and real-time MQTT message bridging via Eclipse Mosquitto broker.", "2. Core Backend Infrastructure: ")

    add_bullet_p(doc, "Integration with Agromonitoring REST API for automatic polygon synchronization, Sentinel-2/Landsat-8 multi-spectral satellite scene acquisition, normalized difference vegetation index (NDVI) mapping, true-color imagery, and historical weather/soil data. Automated AI context generation engine executing Google Gemini 1.5/2.0 LLM with deterministic agronomic fallback rule engines.", "3. Remote Sensing & Generative AI Advisory: ")

    add_bullet_p(doc, "Native SwiftUI mobile application engineered using MVVM and Clean Architecture principles. Features interactive MapKit field polygon boundary drawing, real-time telemetry line charts, multi-spectral NDVI layer rendering, AI recommendation feeds with farmer feedback lifecycle, and an interactive agronomic chatbot supporting photo/image attachments.", "4. Native iOS Mobile Application: ")

    add_heading_2(doc, "1.3 In-Scope vs. Out-of-Scope Boundaries")
    
    # Table for In-Scope / Out-of-Scope
    t_scope = doc.add_table(rows=1, cols=2)
    t_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_scope.rows[0]
    hdr.cells[0].paragraphs[0].text = "In-Scope Functional Features"
    hdr.cells[1].paragraphs[0].text = "Out-of-Scope (Future Roadmaps)"
    style_table_header(hdr)

    scope_data = [
        ("• Firebase Auth user authentication & auto-sync profile", "• Automated physical irrigation hardware valve control"),
        ("• Interactive field polygon creation & area (ha) calculation", "• Drone raw multi-spectral geotiff processing"),
        ("• Real-time ESP32 IoT telemetry ingestion via MQTT & Serial", "• Supply chain financial transaction ledger & marketplace"),
        ("• TimescaleDB hypertable telemetry storage & hourly aggregation", "• Multi-tenant enterprise corporate farm management portal"),
        ("• Satellite NDVI imagery, TrueColor tiles & soil weather sync", "• Offline LLM execution on embedded edge hardware"),
        ("• Autonomous & on-demand Gemini AI agronomic advice loop", "• Automated robotic pesticide spray hardware integration"),
        ("• Farmer feedback tracking (Implemented / Ignored / Pending)", "• Cross-platform Android app (iOS SwiftUI is primary)"),
        ("• Multimodal AI Chatbot with image attachment analysis", "• Physical soil chemical sample laboratory integration")
    ]

    for in_s, out_s in scope_data:
        row = t_scope.add_row()
        row.cells[0].paragraphs[0].text = in_s
        row.cells[1].paragraphs[0].text = out_s

    style_table_body(t_scope, [3.25, 3.25])
    doc.add_paragraph()

    # ==========================================
    # SECTION 2: FUNCTIONAL & NON-FUNCTIONAL REQUIREMENTS
    # ==========================================
    add_heading_1(doc, "2. Functional & Non-Functional Requirements")

    add_heading_2(doc, "2.1 Functional Requirements (FRs)")
    
    t_fr = doc.add_table(rows=1, cols=4)
    t_fr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_fr.rows[0]
    hdr.cells[0].paragraphs[0].text = "Req ID"
    hdr.cells[1].paragraphs[0].text = "Module / Area"
    hdr.cells[2].paragraphs[0].text = "Requirement Title"
    hdr.cells[3].paragraphs[0].text = "Detailed Functional Description"
    style_table_header(hdr)

    fr_data = [
        ("FR-01", "User Authentication", "Firebase Token Auth & Local Sync", "The system shall verify Firebase Bearer ID tokens on protected endpoints and auto-provision local user records in PostgreSQL upon first login."),
        ("FR-02", "Field Mapping", "Polygon Boundary & Area Calculation", "Farmers shall draw field boundaries on an interactive MapKit view. The system shall validate coordinates (min 3 points), format WKT geometries (SRID 4326), compute area in hectares, and sync with Agromonitoring polygon API."),
        ("FR-03", "IoT Pairing", "Sensor Verification & Binding", "The system shall verify sensor node status via device heartbeat (/api/sensors/verify/{id}) and allow binding sensors to specific field polygon IDs."),
        ("FR-04", "Telemetry Ingestion", "MQTT Payload Storage & Aggregation", "The backend MQTT service shall subscribe to agrivision/sensors/+/readings, parse JSON telemetry (temp, moisture, pH, EC, NPK), ingest readings into TimescaleDB hypertables, and trigger hourly bucket aggregations."),
        ("FR-05", "Satellite Remote Sensing", "NDVI & Weather Acquisition", "The background worker shall periodically fetch Sentinel-2 satellite imagery metadata, compute cloud cover %, download NDVI & TrueColor map tiles, and store soil moisture/weather metrics."),
        ("FR-06", "AI Reasoning Engine", "Autonomous Context Advisory", "The AI service shall assemble multi-source field state (sensors + weather + satellite), prompt Gemini LLM with safety constraints, generate actionable advice with confidence scores, and persist in field_recommendations."),
        ("FR-07", "Fallback Rules", "Offline / LLM Outage Fallback", "If Gemini API is unreachable or returns invalid format, the system shall trigger deterministic rule-based agronomic logic to prevent zero recommendation states."),
        ("FR-08", "Recommendation Feedback", "Farmer Outcome Tracking", "Farmers shall view recommendations on the iOS app and log implementation status (Pending, Implemented, Ignored) with optional outcome notes."),
        ("FR-09", "Multimodal AI Chat", "Interactive Agronomist & Image Diagnosis", "Farmers shall send natural language agronomy queries and attach field photos. The backend shall validate image MIME types/SHA256, store files in media storage, and prompt Gemini with image + context data."),
        ("FR-10", "Chat Memory", "Rolling Thread Summarization", "The chat service shall maintain rolling thread summaries (AIChatThread) to compress historical context without exceeding LLM context windows.")
    ]

    for r_id, mod, title, desc in fr_data:
        row = t_fr.add_row()
        row.cells[0].paragraphs[0].text = r_id
        row.cells[1].paragraphs[0].text = mod
        row.cells[2].paragraphs[0].text = title
        row.cells[3].paragraphs[0].text = desc

    style_table_body(t_fr, [0.8, 1.3, 1.6, 2.8])
    doc.add_paragraph()

    add_heading_2(doc, "2.2 Non-Functional Requirements (NFRs)")

    t_nfr = doc.add_table(rows=1, cols=4)
    t_nfr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_nfr.rows[0]
    hdr.cells[0].paragraphs[0].text = "Category"
    hdr.cells[1].paragraphs[0].text = "Metric / Parameter"
    hdr.cells[2].paragraphs[0].text = "Target Standard"
    hdr.cells[3].paragraphs[0].text = "Architectural Implementation Strategy"
    style_table_header(hdr)

    nfr_data = [
        ("NFR-01: Performance", "API Response Latency", "< 500 ms for core GET APIs", "Asynchronous FastAPI async/await handlers, connection pooling with asyncpg & httpx, TTL caching for satellite requests."),
        ("NFR-02: Performance", "Telemetry Ingestion Speed", "< 100 ms per sensor payload", "Mosquitto MQTT broker lightweight pub/sub, non-blocking background queue ingestion."),
        ("NFR-03: Scalability", "Time-Series Query Scaling", "Millions of sensor rows", "TimescaleDB hypertables partitioned by 7-day time chunks; automated continuous aggregates for hourly buckets."),
        ("NFR-04: Scalability", "Geospatial Indexing", "Thousands of fields", "PostGIS GIST spatial indexing on field boundary geometry column (SRID 4326)."),
        ("NFR-05: Reliability", "Service Availability", "99.9% Uptime", "Dockerized container deployment with restart policies; background workers decoupled from API HTTP request handlers."),
        ("NFR-06: Reliability", "Fault Tolerance", "Zero-downtime external outage", "Deterministic agronomic rule engine fallback when external satellite or LLM services fail."),
        ("NFR-07: Security", "Authentication & Authz", "100% Protected Endpoints", "Firebase Admin SDK bearer token verification middleware; strict user-field spatial ownership checks."),
        ("NFR-08: Security", "Data Integrity", "SHA256 File Hashing", "Media attachments validated via MIME type whitelist, file size limits (< 10MB), and unique SHA256 checksums."),
        ("NFR-09: Usability", "UI Rendering Efficiency", "60 FPS Native iOS", "SwiftUI declarative UI architecture, MapKit integration, AsyncImage caching, reactive Combine state management.")
    ]

    for cat, metric, target, impl in nfr_data:
        row = t_nfr.add_row()
        row.cells[0].paragraphs[0].text = cat
        row.cells[1].paragraphs[0].text = metric
        row.cells[2].paragraphs[0].text = target
        row.cells[3].paragraphs[0].text = impl

    style_table_body(t_nfr, [1.3, 1.4, 1.4, 2.4])
    doc.add_paragraph()

    # ==========================================
    # SECTION 3: USE CASE DIAGRAM
    # ==========================================
    add_heading_1(doc, "3. Use Case Diagram")

    add_body_p(doc, "The AgriVision system interacts with five primary external actors: Farmer/Agronomist (User), ESP32 IoT Node (Hardware), Agromonitoring REST API (Satellite Data Provider), Google Gemini API (AI Provider), and Firebase Authentication Service.")

    uc_diagram_text = """+-----------------------------------------------------------------------------------+
|                                 AGRIVISION SYSTEM                                  |
|                                                                                   |
|  +--------------------+                                                           |
|  | Authenticate User  |<------- (1) Verify Token --------> [ Firebase Auth ]      |
|  +--------------------+                                                           |
|           ^                                                                       |
|           |                                                                       |
|  +--------------------+       (2) Sync Polygon                                    |
|  | Create & Map Field |----------------------------------> [ Agromonitoring API ]  |
|  +--------------------+                                                           |
|           ^                                                                       |
|           |                                                                       |
|  +--------------------+       (3) Sensor Reading                                  |
|  | Pair IoT Sensor    |<---------------------------------- [ ESP32 IoT Node ]     |
|  +--------------------+                                                           |
|           ^                                                                       |
| [Farmer / |                                                                       |
|  User]----+-------------------+                                                   |
|           | View Field Dash   |                                                   |
|           +-------------------+                                                   |
|           |                                                                       |
|           |  +---------------------+    (4) Generate Context                      |
|           +->| View Recommendations|---------------------> [ Google Gemini API ]  |
|           |  +---------------------+                                              |
|           |           |                                                           |
|           |           v                                                           |
|           |  +---------------------+                                              |
|           +->| Log Feedback Status |                                              |
|           |  +---------------------+                                              |
|           |                                                                       |
|           |  +---------------------+    (5) Vision + Chat                         |
|           +->| Chat with AI Advisor|---------------------> [ Google Gemini API ]  |
|              +---------------------+                                              |
+-----------------------------------------------------------------------------------+"""
    add_code_block(doc, uc_diagram_text, "Visual Use Case Diagram (Actor & System Interaction)")

    add_heading_2(doc, "3.1 Actor Descriptions")
    add_bullet_p(doc, "Primary end-user who draws field boundaries, views live telemetry/NDVI dashboard, receives AI advice, and logs feedback.", "1. Farmer / Agronomist: ")
    add_bullet_p(doc, "Edge hardware sensing unit that measures soil parameters and streams telemetry over WiFi/MQTT or USB Serial.", "2. ESP32 IoT Sensor Node: ")
    add_bullet_p(doc, "External cloud satellite service providing multi-spectral imagery, polygon sync, and weather metadata.", "3. Agromonitoring API: ")
    add_bullet_p(doc, "Generative AI engine performing multi-modal reasoning over field context and diagnostic photo attachments.", "4. Google Gemini AI Engine: ")
    add_bullet_p(doc, "Identity provider managing user registration, OAuth tokens, and mobile session security.", "5. Firebase Auth: ")

    doc.add_paragraph()

    # ==========================================
    # SECTION 4: ADOPTED METHODOLOGY
    # ==========================================
    add_heading_1(doc, "4. Adopted Methodology")
    add_body_p(doc, "The AgriVision system was developed following the Agile Scrum Methodology, structured into 2-week iterative sprints. This approach enabled rapid prototyping of IoT hardware firmware, backend microservices, and mobile UI features while supporting continuous testing and feedback integration.")

    add_heading_2(doc, "4.1 Sprint Development Lifecycle")
    
    add_bullet_p(doc, "System architecture modeling, database schema definition (PostGIS/TimescaleDB), hardware sensor pin configuration, and API contract design.", "Sprint 1-2 (Architectural Blueprint & Foundations): ")
    add_bullet_p(doc, "ESP32 C++ firmware setup, Modbus RS485 NPK sensor integration, Mosquitto MQTT broker configuration, and Python serial bridge bridge script implementation.", "Sprint 3-4 (IoT Firmware & Telemetry Pipeline): ")
    add_bullet_p(doc, "FastAPI core backend setup, Firebase Auth middleware, spatial polygon endpoints (/api/fields), and TimescaleDB hypertable ingestion pipelines.", "Sprint 5-6 (Backend Microservices & Spatial DB): ")
    add_bullet_p(doc, "Agromonitoring satellite API client, NDVI map tile fetching, Google Gemini AI context engineering, and deterministic fallback rule engine.", "Sprint 7-8 (Satellite Sync & AI Advisory Engine): ")
    add_bullet_p(doc, "Native SwiftUI mobile app development, MapKit boundary drawing tools, Combine reactive state binding, telemetry line charts, and AI recommendation feed.", "Sprint 9-10 (Native iOS Application Development): ")
    add_bullet_p(doc, "End-to-end integration testing, performance optimization, background worker stress testing, security hardening, and documentation assembly.", "Sprint 11-12 (Integration, Testing & QA): ")

    doc.add_paragraph()

    # ==========================================
    # SECTION 5: WORK PLAN (SCHEDULE & WBS)
    # ==========================================
    add_heading_1(doc, "5. Work Plan (MS Project Schedule & WBS)")
    add_body_p(doc, "Below is the formal Work Breakdown Structure (WBS) representing the 12-week development schedule, compatible with MS Project format.")

    t_wbs = doc.add_table(rows=1, cols=6)
    t_wbs.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_wbs.rows[0]
    hdr.cells[0].paragraphs[0].text = "WBS ID"
    hdr.cells[1].paragraphs[0].text = "Task Description"
    hdr.cells[2].paragraphs[0].text = "Duration"
    hdr.cells[3].paragraphs[0].text = "Start Date"
    hdr.cells[4].paragraphs[0].text = "Finish Date"
    hdr.cells[5].paragraphs[0].text = "Predecessors"
    style_table_header(hdr)

    wbs_data = [
        ("1.0", "PROJECT INITIATION & REQUIREMENTS", "10 days", "2026-05-04", "2026-05-15", "-"),
        ("1.1", "System Requirements & Scope Definition", "5 days", "2026-05-04", "2026-05-08", "-"),
        ("1.2", "Architecture Design & Technology Stack Selection", "5 days", "2026-05-11", "2026-05-15", "1.1"),
        ("2.0", "DATABASE & BACKEND INFRASTRUCTURE", "15 days", "2026-05-18", "2026-06-05", "1.2"),
        ("2.1", "PostgreSQL + PostGIS & TimescaleDB Setup", "5 days", "2026-05-18", "2026-05-22", "1.2"),
        ("2.2", "FastAPI Core Setup & Firebase Auth Integration", "5 days", "2026-05-25", "2026-05-29", "2.1"),
        ("2.3", "Spatial Fields API & Polygon Validation Logic", "5 days", "2026-06-01", "2026-06-05", "2.2"),
        ("3.0", "IOT FIRMWARE & TELEMETRY PIPELINE", "10 days", "2026-06-08", "2026-06-19", "2.1"),
        ("3.1", "ESP32 C++ Sensor Drivers (Soil, Temp, NPK)", "5 days", "2026-06-08", "2026-06-12", "1.2"),
        ("3.2", "MQTT Broker & Serial-to-MQTT Bridge Service", "5 days", "2026-06-15", "2026-06-19", "3.1"),
        ("4.0", "SATELLITE & GENERATIVE AI ENGINE", "15 days", "2026-06-22", "2026-07-10", "2.3"),
        ("4.1", "Agromonitoring REST API Client & Satellite Worker", "5 days", "2026-06-22", "2026-06-26", "2.3"),
        ("4.2", "Google Gemini AI Advisory & Fallback Rules Engine", "5 days", "2026-06-29", "2026-07-03", "4.1"),
        ("4.3", "Multimodal AI Chat & Attachment Storage Pipeline", "5 days", "2026-07-06", "2026-07-10", "4.2"),
        ("5.0", "NATIVE IOS MOBILE APPLICATION", "15 days", "2026-07-13", "2026-07-31", "2.3, 3.2, 4.3"),
        ("5.1", "SwiftUI Architecture & MapKit Polygon Editor", "5 days", "2026-07-13", "2026-07-17", "2.3"),
        ("5.2", "Field Dashboard, Telemetry Charts & Satellite Views", "5 days", "2026-07-20", "2026-07-24", "5.1"),
        ("5.3", "AI Recommendation Feed & Chat Interface", "5 days", "2026-07-27", "2026-07-31", "5.2"),
        ("6.0", "TESTING, QA & SYSTEM DEPLOYMENT", "5 days", "2026-07-27", "2026-07-31", "5.3"),
        ("6.1", "End-to-End System QA & Security Hardening", "5 days", "2026-07-27", "2026-07-31", "5.3")
    ]

    for w_id, task, dur, st, fn, pred in wbs_data:
        row = t_wbs.add_row()
        row.cells[0].paragraphs[0].text = w_id
        row.cells[1].paragraphs[0].text = task
        row.cells[2].paragraphs[0].text = dur
        row.cells[3].paragraphs[0].text = st
        row.cells[4].paragraphs[0].text = fn
        row.cells[5].paragraphs[0].text = pred

    style_table_body(t_wbs, [0.6, 2.5, 0.8, 0.9, 0.9, 0.8])
    doc.add_paragraph()

    # ==========================================
    # SECTION 6: ENTITY RELATIONSHIP DIAGRAM (ERD)
    # ==========================================
    add_heading_1(doc, "6. Entity Relationship Diagram (ERD)")
    add_body_p(doc, "The AgriVision database is built on PostgreSQL with PostGIS extensions for spatial polygons and TimescaleDB extensions for time-series hypertables. Below is the structural Entity Relationship Diagram.")

    erd_text = """[ USERS ] (1) ───< (N) [ FIELDS ] (1) ───< (N) [ SENSORS ]
  │                       │                         │
  │ (1)                   │ (1)                     │ (1)
  v                       v                         v
[ SENSOR_READINGS ]     [ SATELLITE_SCENES ]      [ SENSOR_READINGS ] (Hypertable)
                          │                         │
                          │ (1)                     │ (1)
                          v                         v
                        [ AI_ANALYSIS_RUNS ]      [ SENSOR_READINGS_HOURLY ]
                          │
                          │ (1)
                          v
                        [ FIELD_RECOMMENDATIONS ]
                          │
                          │ (1)
                          v
                        [ AI_CHAT_THREADS ] (1) ───< (N) [ AI_CHAT_MESSAGES ] (1) ───< (N) [ CHAT_ATTACHMENTS ]"""
    add_code_block(doc, erd_text, "Database ERD Entity Connections & Cardinalities")

    add_heading_2(doc, "6.1 Detailed Table Cardinalities")
    add_bullet_p(doc, "One User can register multiple Fields. A Field belongs to exactly one User owner.", "User -> Fields: One-to-Many (1:N) ")
    add_bullet_p(doc, "A Field can have multiple bound Sensors. A Sensor can be bound to at most one Field (ON DELETE SET NULL).", "Field -> Sensors: One-to-Many (1:N) ")
    add_bullet_p(doc, "A Sensor emits multiple high-frequency SensorReadings stored in TimescaleDB hypertable (ON DELETE CASCADE).", "Sensor -> SensorReadings: One-to-Many (1:N) ")
    add_bullet_p(doc, "A Field retains historical multi-spectral SatelliteScenes acquired from Agromonitoring.", "Field -> SatelliteScenes: One-to-Many (1:N) ")
    add_bullet_p(doc, "Each AI evaluation run generates multiple actionable FieldRecommendations.", "AIAnalysisRun -> FieldRecommendations: One-to-Many (1:N) ")
    add_bullet_p(doc, "Each Field maintains exactly one active AIChatThread (1:1) containing multiple AIChatMessages (1:N) and ChatAttachments (1:N).", "Field -> AIChatThread -> AIChatMessages: One-to-Many (1:N) ")

    doc.add_paragraph()

    # ==========================================
    # SECTION 7: ARCHITECTURE DESIGN DIAGRAM
    # ==========================================
    add_heading_1(doc, "7. Architecture Design Diagram")
    add_body_p(doc, "AgriVision employs a multi-tiered, decoupled microservices architecture designed for high throughput, low latency, and high reliability.")

    arch_text = """+-----------------------------------------------------------------------------------+
|                            PRESENTATION LAYER (NATIVE IOS)                        |
|   SwiftUI Views | MapKit Boundary Drawing | Charts | MVVM ViewModels | Combine State  |
+-----------------------------------------------------------------------------------+
                                          │  HTTP / REST (Firebase Bearer Token)
                                          v
+-----------------------------------------------------------------------------------+
|                         APPLICATION & API GATEWAY LAYER                           |
|   FastAPI ASGI Server (Uvicorn) | Firebase Admin Auth | Async Routers & Pydantic   |
+-----------------------------------------------------------------------------------+
       │                                  │                                  │
       │ Internal Async Call              │ MQTT Pub/Sub                     │ Background Loop
       v                                  v                                  v
+-----------------------+      +-----------------------+      +-----------------------+
| SATELLITE SYNC WORKER |      | MOSQUITTO MQTT BROKER |      | AI REASONING WORKER   |
| (Hourly Agromonitoring|      | agrivision/sensors/   |      | (4-Hour Gemini LLM    |
| Polygon & NDVI Sync)  |      | +/readings            |      | Advisor & Fallback)   |
+-----------------------+      +-----------------------+      +-----------------------+
       │                                  │                                  │
       v                                  v                                  v
+-----------------------------------------------------------------------------------+
|                            PERSISTENCE & DATA STORAGE                             |
|   PostgreSQL + PostGIS (Spatial Polygons) | TimescaleDB Hypertables (Telemetry)   |
+-----------------------------------------------------------------------------------+
                                          ^
                                          │ USB Serial Bridge / WiFi MQTT
+-----------------------------------------------------------------------------------+
|                            HARDWARE & IOT EDGE LAYER                              |
|   ESP32 Microcontroller | Modbus RS485 NPK | Soil Moisture & DS18B20 Temp Sensors  |
+-----------------------------------------------------------------------------------+"""
    add_code_block(doc, arch_text, "AgriVision Multi-Tier Architectural Diagram")

    add_heading_2(doc, "7.1 Subsystem Responsibilities")
    add_bullet_p(doc, "Provides farmer-facing interactive maps, live telemetry visualization, AI recommendation feedback controls, and chat with photo upload.", "Presentation Layer (iOS): ")
    add_bullet_p(doc, "Enforces Firebase token authorization, executes geospatial boundary validation, handles file uploads with SHA256 checksums, and exposes REST endpoints.", "API Gateway (FastAPI): ")
    add_bullet_p(doc, "Executes decoupled background event loops: hourly satellite imagery syncs and 4-hour AI reasoning batch runs without blocking HTTP API requests.", "Background Worker Layer: ")
    add_bullet_p(doc, "Stores relational models in PostgreSQL, geospatial geometries in PostGIS (SRID 4326), and telemetry readings in TimescaleDB hypertables.", "Persistence Layer: ")
    add_bullet_p(doc, "ESP32 hardware node collects soil moisture, temperature, and NPK metrics every 30 seconds and publishes to MQTT broker.", "IoT Layer: ")

    doc.add_paragraph()

    # ==========================================
    # SECTION 8: SEQUENCE DIAGRAMS / USAGE SCENARIOS
    # ==========================================
    add_heading_1(doc, "8. Sequence Diagrams / Usage Scenarios")

    add_heading_2(doc, "8.1 Scenario 1: Field Creation & Satellite Polygon Sync")
    seq1_text = """Farmer (iOS)              FastAPI Backend            PostgreSQL / PostGIS         Agromonitoring API
    │                              │                               │                          │
    ├─ 1. POST /api/fields ───────>│                               │                          │
    │   (Name, Boundary Coords)    ├─ 2. Validate WKT Geometry ───>│                          │
    │                              ├─ 3. Save Field Record ───────>│                          │
    │                              │                               │                          │
    │                              ├─ 4. POST /polygons ─────────────────────────────────────>│
    │                              │   (GeoJSON Polygon)                                      │
    │                              │<── 5. Return poly_id ────────────────────────────────────┤
    │                              │                                                          │
    │                              ├─ 6. Update Field agromonitory_poly_id ──────────────────>│
    │<─ 7. 201 Created (Field JSON)│                                                          │"""
    add_code_block(doc, seq1_text, "Sequence Diagram 1: Field Creation & Satellite Sync")

    add_heading_2(doc, "8.2 Scenario 2: Real-Time IoT Telemetry Ingestion")
    seq2_text = """ESP32 Node                Mosquitto MQTT            FastAPI MQTT Service         TimescaleDB
    │                              │                               │                          │
    ├─ 1. Publish Telemetry ──────>│                               │                          │
    │   agrivision/sensors/...     ├─ 2. Forward Payload ─────────>│                          │
    │                              │                               ├─ 3. Parse JSON Telemetry │
    │                              │                               ├─ 4. INSERT SensorReading─>│
    │                              │                               │   (Time, Moisture, NPK)  │
    │                              │                               ├─ 5. Update last_seen ───>│"""
    add_code_block(doc, seq2_text, "Sequence Diagram 2: IoT Telemetry Ingestion Pipeline")

    add_heading_2(doc, "8.3 Scenario 3: AI Agronomy Recommendation Loop")
    seq3_text = """Scheduler Worker           AgriVision Database        Gemini AI Engine             iOS App
    │                              │                               │                          │
    ├─ 1. Trigger (Every 4h)       │                               │                          │
    ├─ 2. Fetch Field Context ────>│                               │                          │
    │   (Sensors, Weather, NDVI)   │                               │                          │
    ├─ 3. Build AI Prompt ─────────┼──────────────────────────────>│                          │
    │                              │<── 4. Return Advice JSON ─────┤                          │
    │                              │    (Category, Priority, Rec)  │                          │
    ├─ 5. Save Recommendation ────>│                               │                          │
    │                              │<── 6. GET /recommendations ──────────────────────────────┤
    │                              │─── 7. Return AI Recs Feed ──────────────────────────────>│"""
    add_code_block(doc, seq3_text, "Sequence Diagram 3: AI Recommendation Loop")

    add_heading_2(doc, "8.4 Scenario 4: Interactive Multimodal AI Chat")
    seq4_text = """Farmer (iOS)              FastAPI Chat Router         Media Storage              Gemini AI Engine
    │                              │                               │                          │
    ├─ 1. POST /chat ─────────────>│                               │                          │
    │   (Text Query + Photo)       ├─ 2. Save Attachment (SHA256)─>│                          │
    │                              ├─ 3. Save User Message         │                          │
    │                              ├─ 4. Compile Memory & Image ──┼─────────────────────────>│
    │                              │<── 5. AI Agronomist Reply ────┼──────────────────────────┤
    │                              ├─ 6. Save AI Response          │                          │
    │<─ 7. 200 OK (AI Response JSON)│                              │                          │"""
    add_code_block(doc, seq4_text, "Sequence Diagram 4: Multimodal AI Chat Scenario")

    doc.add_paragraph()

    # ==========================================
    # SECTION 9: CLASS DIAGRAM
    # ==========================================
    add_heading_1(doc, "9. Class Diagram")

    add_heading_2(doc, "9.1 Backend Architecture Class Structure")
    cls_backend_text = """+------------------------------------+          +------------------------------------+
|               User                 |          |               Field                |
+------------------------------------+          +------------------------------------+
| + id: UUID                         | 1      N | + id: UUID                         |
| + firebase_uid: String             |----------| + owner_id: UUID                   |
| + email: String                    |          | + name: String                     |
| + created_at: DateTime             |          | + boundary: Geometry(Polygon,4326) |
+------------------------------------+          | + area_ha: Float                   |
                                                | + agromonitory_poly_id: String     |
                                                | + latest_ndvi: Float               |
                                                +------------------------------------+
                                                                  │ 1
                                                                  │
                                                                  │ N
+------------------------------------+          +------------------------------------+
|        FieldRecommendation         |          |               Sensor               |
+------------------------------------+          +------------------------------------+
| + id: UUID                         |          | + id: UUID                         |
| + field_id: UUID                   |          | + field_id: UUID                   |
| + category: String                 |          | + device_id: String                |
| + priority: String                 |          | + sensor_type: String              |
| + advice: Text                     |          | + battery_level: Float             |
| + confidence: Float                |          | + last_seen: DateTime              |
| + status: String                   |          +------------------------------------+
+------------------------------------+"""
    add_code_block(doc, cls_backend_text, "Backend Core Model Class Diagram")

    add_heading_2(doc, "9.2 iOS MVVM Application Class Structure")
    cls_ios_text = """+-----------------------------------+        +-----------------------------------+
|      FieldDashboardViewModel      |        |          AgriDataService          |
+-----------------------------------+        +-----------------------------------+
| - dataService: AgriDataService    |------->| + fetchFields(): [Field]          |
| @Published fields: [Field]        |        | + fetchRecommendations(id): [Rec] |
| @Published recs: [Recommendation] |        | + createField(FieldRequest): Field|
| + loadDashboardData()             |        | + sendChatMessage(msg, img): Chat |
| + updateFeedback(recId, status)   |        +-----------------------------------+
+-----------------------------------+                          │
                                                               v
+-----------------------------------+        +-----------------------------------+
|         AIChatViewModel           |        |            APIConstants           |
+-----------------------------------+        +-----------------------------------+
| - chatService: AgriDataService    |        | + baseURL: String = "localhost"   |
| @Published messages: [ChatMessage]|        | + fieldURL(id): URL               |
| + sendMessage(text, photo)        |        | + chatURL(id): URL                |
+-----------------------------------+        +-----------------------------------+"""
    add_code_block(doc, cls_ios_text, "iOS MVVM & Service Layer Class Diagram")

    doc.add_paragraph()

    # ==========================================
    # SECTION 10: DATABASE DESIGN
    # ==========================================
    add_heading_1(doc, "10. Database Design & Data Dictionary")
    add_body_p(doc, "The AgriVision relational schema is designed for PostGIS spatial operations and TimescaleDB time-series hypertable efficiency. Below is the complete data dictionary.")

    db_tables = [
        ("1. users", "Stores registered user accounts synchronized with Firebase Auth.", [
            ("id", "UUID", "PRIMARY KEY", "Unique user identifier (UUIDv4)"),
            ("firebase_uid", "VARCHAR(128)", "UNIQUE, NOT NULL", "Firebase Authentication UID"),
            ("email", "VARCHAR(320)", "UNIQUE", "User email address"),
            ("created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Account creation timestamp")
        ]),
        ("2. fields", "Stores geospatial field boundaries, crop metadata, and satellite IDs.", [
            ("id", "UUID", "PRIMARY KEY", "Unique field identifier"),
            ("owner_id", "UUID", "FK (users.id)", "Field owner user ID"),
            ("name", "VARCHAR(100)", "NOT NULL", "Farmer defined field name"),
            ("crop_type", "VARCHAR(80)", "NULLABLE", "Crop type (e.g., Wheat, Cotton)"),
            ("plantation_date", "TIMESTAMPTZ", "NULLABLE", "Planting timestamp"),
            ("boundary", "GEOMETRY(POLYGON,4326)", "NOT NULL", "PostGIS spatial polygon (EPSG:4326)"),
            ("area_ha", "FLOAT", "NOT NULL", "Calculated area in hectares"),
            ("agromonitory_poly_id", "VARCHAR(64)", "NULLABLE", "External Agromonitoring polygon ID"),
            ("latest_ndvi", "FLOAT", "NULLABLE", "Latest vegetation index score"),
            ("status", "VARCHAR(20)", "DEFAULT 'active'", "Field status (active/archived)")
        ]),
        ("3. sensors", "Stores IoT hardware sensor node metadata and status.", [
            ("id", "UUID", "PRIMARY KEY", "Unique sensor system ID"),
            ("field_id", "UUID", "FK (fields.id)", "Bound field ID (ON DELETE SET NULL)"),
            ("device_id", "VARCHAR(100)", "UNIQUE, NOT NULL", "Hardware device serial ID"),
            ("sensor_type", "VARCHAR(50)", "DEFAULT 'multi_sensor'", "Sensor capability type"),
            ("battery_level", "FLOAT", "NULLABLE", "Battery percentage (0-100%)"),
            ("last_seen", "TIMESTAMPTZ", "NULLABLE", "Last heartbeat timestamp")
        ]),
        ("4. sensor_readings (TimescaleDB Hypertable)", "Stores high-frequency sensor telemetry time-series.", [
            ("time", "TIMESTAMPTZ", "PRIMARY KEY (Composite)", "Telemetry measurement timestamp"),
            ("sensor_id", "UUID", "PRIMARY KEY, FK (sensors.id)", "Sensor node ID"),
            ("temperature", "FLOAT", "NULLABLE", "Soil / Air temperature (°C)"),
            ("moisture", "FLOAT", "NULLABLE", "Soil moisture percentage (%)"),
            ("humidity", "FLOAT", "NULLABLE", "Relative humidity (%)"),
            ("ph", "FLOAT", "NULLABLE", "Soil pH level"),
            ("ec", "FLOAT", "NULLABLE", "Electrical conductivity (dS/m)"),
            ("npk_n", "FLOAT", "NULLABLE", "Nitrogen content (mg/kg)"),
            ("npk_p", "FLOAT", "NULLABLE", "Phosphorus content (mg/kg)"),
            ("npk_k", "FLOAT", "NULLABLE", "Potassium content (mg/kg)")
        ]),
        ("5. field_recommendations", "Stores generated Generative AI agronomic recommendations.", [
            ("id", "UUID", "PRIMARY KEY", "Recommendation identifier"),
            ("field_id", "UUID", "FK (fields.id)", "Target field ID"),
            ("category", "VARCHAR(50)", "NOT NULL", "Category (Irrigation, Fertilizer, Pest)"),
            ("priority", "VARCHAR(20)", "DEFAULT 'medium'", "Priority (high, medium, low)"),
            ("advice", "TEXT", "NOT NULL", "Actionable advice text"),
            ("confidence", "FLOAT", "NULLABLE", "AI model confidence score (0.0 - 1.0)"),
            ("status", "VARCHAR(20)", "DEFAULT 'pending'", "Status (pending, implemented, ignored)")
        ]),
        ("6. ai_chat_messages", "Stores multi-turn AI agronomy chatbot messages.", [
            ("id", "UUID", "PRIMARY KEY", "Chat message ID"),
            ("field_id", "UUID", "FK (fields.id)", "Associated field ID"),
            ("role", "VARCHAR(20)", "NOT NULL", "Message role (user / assistant)"),
            ("content", "TEXT", "NOT NULL", "Chat message text content"),
            ("created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Message timestamp")
        ]),
        ("7. chat_attachments", "Stores diagnostic photo attachments sent in AI chat.", [
            ("id", "UUID", "PRIMARY KEY", "Attachment ID"),
            ("field_id", "UUID", "FK (fields.id)", "Associated field ID"),
            ("message_id", "UUID", "FK (ai_chat_messages.id)", "Associated message ID"),
            ("storage_key", "VARCHAR(500)", "UNIQUE, NOT NULL", "File path / media key"),
            ("mime_type", "VARCHAR(50)", "NOT NULL", "MIME type (image/jpeg, image/png)"),
            ("sha256", "VARCHAR(64)", "NOT NULL", "SHA256 checksum for integrity")
        ])
    ]

    for t_title, t_desc, cols in db_tables:
        add_heading_2(doc, f"10.{t_title}")
        add_body_p(doc, t_desc)
        
        t_dict = doc.add_table(rows=1, cols=4)
        t_dict.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t_dict.rows[0]
        hdr.cells[0].paragraphs[0].text = "Column Name"
        hdr.cells[1].paragraphs[0].text = "Data Type"
        hdr.cells[2].paragraphs[0].text = "Constraints"
        hdr.cells[3].paragraphs[0].text = "Description & Business Logic"
        style_table_header(hdr)

        for col_name, dtype, cons, cdesc in cols:
            row = t_dict.add_row()
            row.cells[0].paragraphs[0].text = col_name
            row.cells[1].paragraphs[0].text = dtype
            row.cells[2].paragraphs[0].text = cons
            row.cells[3].paragraphs[0].text = cdesc

        style_table_body(t_dict, [1.5, 1.5, 1.3, 2.2])
        doc.add_paragraph()

    # Output file
    out_dir = "/Users/ahmad/AgriVision/Docs"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "AgriVision_Project_Documentation.docx")
    doc.save(out_path)
    print(f"Document successfully created at: {out_path}")

if __name__ == "__main__":
    generate_document()
